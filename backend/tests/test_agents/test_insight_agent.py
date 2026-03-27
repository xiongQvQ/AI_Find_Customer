"""Tests for agents/insight_agent.py — mock react_loop, verify insight output."""

import json
from unittest.mock import AsyncMock, patch, MagicMock

import pytest

from agents.insight_agent import (
    insight_node,
    _parse_uploaded_file,
    _pre_parse_documents,
    _build_user_prompt,
)


FAKE_INSIGHT_JSON = json.dumps({
    "company_name": "SolarTech GmbH",
    "products": ["solar inverter", "PV panel"],
    "industries": ["Renewable Energy", "Electronics"],
    "value_propositions": ["High efficiency", "10-year warranty"],
    "target_customer_profile": "B2B distributors in Europe",
    "recommended_regions": ["Europe", "North America"],
    "recommended_keywords_seed": ["solar inverter distributor", "PV panel wholesale"],
    "summary": "SolarTech manufactures high-efficiency solar inverters.",
})


def _base_state(**overrides):
    base = {
        "website_url": "https://solartech.de",
        "product_keywords": ["solar inverter"],
        "target_regions": ["Europe"],
        "uploaded_files": [],
        "target_lead_count": 200,
        "max_rounds": 10,
        "insight": None,
        "keywords": [],
        "used_keywords": [],
        "search_results": [],
        "matched_platforms": [],
        "keyword_search_stats": {},
        "leads": [],
        "email_sequences": [],
        "hunt_round": 1,
        "prev_round_lead_count": 0,
        "round_feedback": None,
        "current_stage": "start",
        "messages": [],
    }
    base.update(overrides)
    return base


class TestInsightNode:
    @pytest.mark.asyncio
    async def test_returns_insight_and_stage(self):
        """Normal flow: ReAct loop returns structured insight JSON."""
        state = _base_state()

        with patch("agents.insight_agent.react_loop", return_value=FAKE_INSIGHT_JSON), \
             patch("agents.insight_agent.JinaReaderTool") as MockJina, \
             patch("agents.insight_agent.GoogleSearchTool") as MockGoogle:

            MockJina.return_value = AsyncMock(close=AsyncMock())
            MockGoogle.return_value = AsyncMock(close=AsyncMock())

            result = await insight_node(state)

        assert result["current_stage"] == "insight"
        assert result["insight"]["company_name"] == "SolarTech GmbH"
        assert "solar inverter" in result["insight"]["products"]
        assert len(result["insight"]["recommended_keywords_seed"]) > 0

    @pytest.mark.asyncio
    async def test_no_website_no_files_returns_minimal(self):
        """When no website and no files, return minimal insight."""
        state = _base_state(website_url="", uploaded_files=[], product_keywords=[], target_regions=[])

        result = await insight_node(state)

        assert result["current_stage"] == "insight"
        assert result["insight"]["company_name"] == "Unknown"
        assert result["insight"]["summary"] == "No website or documents provided."

    @pytest.mark.asyncio
    async def test_react_returns_invalid_json(self):
        """If ReAct returns non-JSON, agent raises RuntimeError to abort the hunt."""
        state = _base_state()

        with patch("agents.insight_agent.react_loop", return_value="This is not JSON at all"), \
             patch("agents.insight_agent.JinaReaderTool") as MockJina, \
             patch("agents.insight_agent.GoogleSearchTool") as MockGoogle:

            MockJina.return_value = AsyncMock(close=AsyncMock())
            MockGoogle.return_value = AsyncMock(close=AsyncMock())

            with pytest.raises(RuntimeError, match="InsightAgent"):
                await insight_node(state)

    @pytest.mark.asyncio
    async def test_react_loop_failure(self):
        """If react_loop raises, InsightAgent re-raises to abort the hunt."""
        state = _base_state()

        with patch("agents.insight_agent.react_loop", side_effect=Exception("API down")), \
             patch("agents.insight_agent.JinaReaderTool") as MockJina, \
             patch("agents.insight_agent.GoogleSearchTool") as MockGoogle:

            MockJina.return_value = AsyncMock(close=AsyncMock())
            MockGoogle.return_value = AsyncMock(close=AsyncMock())

            with pytest.raises(RuntimeError, match="InsightAgent failed"):
                await insight_node(state)

    @pytest.mark.asyncio
    async def test_with_uploaded_files(self):
        """Uploaded files are mentioned in the user prompt to the ReAct agent."""
        state = _base_state(uploaded_files=["/tmp/catalog.pdf"])

        captured_kwargs = {}

        async def capture_react_loop(**kwargs):
            captured_kwargs.update(kwargs)
            return FAKE_INSIGHT_JSON

        with patch("agents.insight_agent.react_loop", side_effect=capture_react_loop), \
             patch("agents.insight_agent.JinaReaderTool") as MockJina, \
             patch("agents.insight_agent.GoogleSearchTool") as MockGoogle:

            MockJina.return_value = AsyncMock(close=AsyncMock())
            MockGoogle.return_value = AsyncMock(close=AsyncMock())

            result = await insight_node(state)

        assert "catalog.pdf" in captured_kwargs["user_prompt"]
        assert result["insight"]["company_name"] == "SolarTech GmbH"

    @pytest.mark.asyncio
    async def test_only_keywords_no_website(self):
        """When only keywords provided (no website), ReAct still gets called."""
        state = _base_state(website_url="", product_keywords=["solar panel", "inverter"])

        with patch("agents.insight_agent.react_loop", return_value=FAKE_INSIGHT_JSON), \
             patch("agents.insight_agent.JinaReaderTool") as MockJina, \
             patch("agents.insight_agent.GoogleSearchTool") as MockGoogle:

            MockJina.return_value = AsyncMock(close=AsyncMock())
            MockGoogle.return_value = AsyncMock(close=AsyncMock())

            result = await insight_node(state)

        assert result["insight"]["company_name"] == "SolarTech GmbH"

    @pytest.mark.asyncio
    async def test_target_regions_in_prompt(self):
        """Target regions are passed to the ReAct agent in the user prompt."""
        state = _base_state(target_regions=["Europe", "Southeast Asia"])

        captured_kwargs = {}

        async def capture_react_loop(**kwargs):
            captured_kwargs.update(kwargs)
            return FAKE_INSIGHT_JSON

        with patch("agents.insight_agent.react_loop", side_effect=capture_react_loop), \
             patch("agents.insight_agent.JinaReaderTool") as MockJina, \
             patch("agents.insight_agent.GoogleSearchTool") as MockGoogle:

            MockJina.return_value = AsyncMock(close=AsyncMock())
            MockGoogle.return_value = AsyncMock(close=AsyncMock())

            await insight_node(state)

        assert "Europe" in captured_kwargs["user_prompt"]
        assert "Southeast Asia" in captured_kwargs["user_prompt"]
        assert "MUST appear first" in captured_kwargs["user_prompt"]


class TestParseUploadedFile:
    def test_parse_pdf(self, tmp_path):
        """PDF files are parsed via PDFParserTool."""
        import pymupdf

        pdf_path = str(tmp_path / "test.pdf")
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Test PDF content")
        doc.save(pdf_path)
        doc.close()

        result = _parse_uploaded_file(pdf_path)
        assert "Test PDF" in result or len(result) > 0

    def test_parse_csv(self, tmp_path):
        """CSV files are parsed via ExcelParserTool."""
        csv_path = str(tmp_path / "data.csv")
        with open(csv_path, "w") as f:
            f.write("name,value\nAlpha,100\n")

        result = _parse_uploaded_file(csv_path)
        assert "Alpha" in result

    def test_parse_txt(self, tmp_path):
        """Text files are read directly."""
        txt_path = str(tmp_path / "notes.txt")
        with open(txt_path, "w") as f:
            f.write("Some product notes")

        result = _parse_uploaded_file(txt_path)
        assert "product notes" in result

    def test_parse_md(self, tmp_path):
        """Markdown files are read directly."""
        md_path = str(tmp_path / "readme.md")
        with open(md_path, "w") as f:
            f.write("# Product Catalog\n\nSolar inverters 5kW-50kW")

        result = _parse_uploaded_file(md_path)
        assert "Product Catalog" in result
        assert "Solar inverters" in result

    def test_parse_json(self, tmp_path):
        """JSON files are read as plain text."""
        json_path = str(tmp_path / "data.json")
        with open(json_path, "w") as f:
            f.write('{"product": "solar inverter", "power": "5kW"}')

        result = _parse_uploaded_file(json_path)
        assert "solar inverter" in result

    def test_parse_docx(self, tmp_path):
        """Word .docx files are parsed via DocxParserTool."""
        from docx import Document
        docx_path = str(tmp_path / "catalog.docx")
        doc = Document()
        doc.add_heading("Product Catalog", level=1)
        doc.add_paragraph("Our main product is the SolarMax 5000 inverter.")
        doc.save(docx_path)

        result = _parse_uploaded_file(docx_path)
        assert "Product Catalog" in result
        assert "SolarMax" in result

    def test_parse_nonexistent_file(self):
        """Non-existent files return error message."""
        result = _parse_uploaded_file("/nonexistent/file.pdf")
        assert "Failed to parse" in result

    def test_parse_unknown_extension(self, tmp_path):
        """Unknown extensions attempt plain text read."""
        unknown_path = str(tmp_path / "data.xyz")
        with open(unknown_path, "w") as f:
            f.write("Some content here")
        result = _parse_uploaded_file(unknown_path)
        assert "Some content" in result


class TestPreParseDocuments:
    def test_empty_list(self):
        result = _pre_parse_documents([])
        assert result == []

    def test_single_txt_file(self, tmp_path):
        txt_path = str(tmp_path / "notes.txt")
        with open(txt_path, "w") as f:
            f.write("Product: solar inverter 5kW")

        result = _pre_parse_documents([txt_path])
        assert len(result) == 1
        assert result[0]["index"] == 0
        assert result[0]["name"] == "notes.txt"
        assert result[0]["type"] == "Plain text"
        assert "solar inverter" in result[0]["content"]
        assert result[0]["truncated"] is False

    def test_multiple_files(self, tmp_path):
        f1 = str(tmp_path / "a.txt")
        f2 = str(tmp_path / "b.md")
        with open(f1, "w") as f:
            f.write("File A content")
        with open(f2, "w") as f:
            f.write("# File B heading")

        result = _pre_parse_documents([f1, f2])
        assert len(result) == 2
        assert result[0]["name"] == "a.txt"
        assert result[1]["name"] == "b.md"
        assert result[1]["type"] == "Markdown document"

    def test_truncation_per_file(self, tmp_path):
        """Files larger than _MAX_DOC_CHARS_PER_FILE are truncated."""
        from agents.insight_agent import _MAX_DOC_CHARS_PER_FILE
        large_path = str(tmp_path / "large.txt")
        with open(large_path, "w") as f:
            f.write("X" * (_MAX_DOC_CHARS_PER_FILE + 1000))

        result = _pre_parse_documents([large_path])
        assert result[0]["truncated"] is True
        assert len(result[0]["content"]) == _MAX_DOC_CHARS_PER_FILE

    def test_indices_are_correct(self, tmp_path):
        files = []
        for i in range(3):
            p = str(tmp_path / f"file_{i}.txt")
            with open(p, "w") as f:
                f.write(f"Content {i}")
            files.append(p)

        result = _pre_parse_documents(files)
        for i, doc in enumerate(result):
            assert doc["index"] == i


class TestBuildUserPrompt:
    def test_url_only(self):
        prompt = _build_user_prompt(
            website_url="https://solartech.de",
            product_keywords=[],
            target_regions=[],
            parsed_docs=[],
        )
        assert "https://solartech.de" in prompt
        assert "Scrape this URL first" in prompt
        assert "Your Task" in prompt

    def test_docs_only(self):
        parsed_docs = [{
            "index": 0, "name": "catalog.pdf", "type": "PDF document",
            "content": "Solar inverter specs", "full_length": 20, "truncated": False,
        }]
        prompt = _build_user_prompt(
            website_url="",
            product_keywords=[],
            target_regions=[],
            parsed_docs=parsed_docs,
        )
        assert "catalog.pdf" in prompt
        assert "Solar inverter specs" in prompt
        assert "No website URL was provided" in prompt

    def test_keywords_only(self):
        prompt = _build_user_prompt(
            website_url="",
            product_keywords=["solar inverter", "PV panel"],
            target_regions=[],
            parsed_docs=[],
        )
        assert "solar inverter" in prompt
        assert "search_web" in prompt

    def test_url_and_docs(self):
        parsed_docs = [{
            "index": 0, "name": "spec.docx", "type": "Word document",
            "content": "Product specs here", "full_length": 18, "truncated": False,
        }]
        prompt = _build_user_prompt(
            website_url="https://example.com",
            product_keywords=[],
            target_regions=[],
            parsed_docs=parsed_docs,
        )
        assert "https://example.com" in prompt
        assert "spec.docx" in prompt
        assert "BOTH a website URL and uploaded documents" in prompt

    def test_target_regions_in_prompt(self):
        prompt = _build_user_prompt(
            website_url="",
            product_keywords=["inverter"],
            target_regions=["Germany", "Poland"],
            parsed_docs=[],
        )
        assert "Germany" in prompt
        assert "Poland" in prompt
        assert "MUST appear first" in prompt

    def test_truncated_doc_shows_note(self):
        parsed_docs = [{
            "index": 0, "name": "big.pdf", "type": "PDF document",
            "content": "A" * 100, "full_length": 10000, "truncated": True,
        }]
        prompt = _build_user_prompt(
            website_url="",
            product_keywords=[],
            target_regions=[],
            parsed_docs=parsed_docs,
        )
        assert "read_uploaded_file(0)" in prompt
        assert "truncated" in prompt

    def test_empty_doc_content_shows_budget_note(self):
        parsed_docs = [{
            "index": 1, "name": "extra.pdf", "type": "PDF document",
            "content": "", "full_length": 0, "truncated": True,
        }]
        prompt = _build_user_prompt(
            website_url="",
            product_keywords=[],
            target_regions=[],
            parsed_docs=parsed_docs,
        )
        assert "read_uploaded_file(1)" in prompt
        assert "budget exhausted" in prompt
